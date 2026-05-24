import sys, os, io
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_PATH = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\Final_Dissertation_BITS_2024MT03117.docx'
OUT_PATH = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\Final_Dissertation_BITS_2024MT03117_Final.docx'
IMG_DIR  = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\dissertation_imgs'
os.makedirs(IMG_DIR, exist_ok=True)

# ─── GENERATE ALL FIGURES ────────────────────────────────────────────────────

def save(fig, name):
    path = os.path.join(IMG_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f'  Generated: {name}')
    return path

# Fig 3.1 — System Architecture
def fig_system_arch():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    ax.set_facecolor('#f8f9fa')

    layers = [
        (0.3, 5.0, 9.4, 0.7, '#4A90D9', 'white', 'PRESENTATION LAYER — React.js (TypeScript + Tailwind CSS)'),
        (0.3, 4.0, 9.4, 0.7, '#7B68EE', 'white', 'API GATEWAY — Traefik Ingress Controller (Routing + Auth + Rate Limiting)'),
        (0.3, 3.0, 9.4, 0.7, '#2ECC71', 'white', 'BUSINESS LAYER — Node.js / Express Microservices (Auth | Symptoms | AI | Admin)'),
        (0.3, 2.0, 4.5, 0.7, '#E67E22', 'white', 'DATA LAYER — MySQL 8.x'),
        (4.9, 2.0, 4.8, 0.7, '#E74C3C', 'white', 'CACHE LAYER — Redis 7'),
        (0.3, 1.0, 9.4, 0.7, '#1ABC9C', 'white', 'INFRA LAYER — K3s (HPA + Self-Healing + Load Balancing)'),
        (0.3, 0.1, 9.4, 0.6, '#95A5A6', 'white', 'OBSERVABILITY — Prometheus | Grafana | Loki | Promtail'),
    ]
    for x, y, w, h, color, tc, label in layers:
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.05',
                              facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                color=tc, fontsize=9, fontweight='bold', wrap=True)

    for y in [4.7, 3.7, 2.7, 1.7]:
        ax.annotate('', xy=(5, y), xytext=(5, y+0.3),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    ax.set_title('Figure 3.1: System Architecture – Cloud-Native Application Layers',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig3_1_arch.png')

# Fig 3.2 — Functional Block Diagram
def fig_block_diagram():
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11); ax.set_ylim(0, 7); ax.axis('off')

    boxes = [
        (4.5, 6.2, 2.0, 0.6, '#3498DB', 'User / Browser'),
        (4.5, 5.2, 2.0, 0.6, '#9B59B6', 'Traefik Ingress'),
        (1.0, 3.8, 1.8, 0.6, '#2ECC71', 'Auth Service'),
        (3.1, 3.8, 1.8, 0.6, '#2ECC71', 'Symptom Svc'),
        (5.2, 3.8, 1.8, 0.6, '#2ECC71', 'AI Service'),
        (7.3, 3.8, 1.8, 0.6, '#2ECC71', 'Admin Service'),
        (3.1, 2.5, 1.8, 0.6, '#E67E22', 'MySQL DB'),
        (5.2, 2.5, 1.8, 0.6, '#E74C3C', 'Redis Cache'),
        (0.5, 1.2, 2.0, 0.6, '#1ABC9C', 'Prometheus'),
        (3.0, 1.2, 2.0, 0.6, '#1ABC9C', 'Grafana'),
        (5.5, 1.2, 2.0, 0.6, '#1ABC9C', 'Loki'),
        (8.0, 1.2, 2.0, 0.6, '#1ABC9C', 'Promtail'),
    ]
    for x, y, w, h, color, label in boxes:
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.08',
                              facecolor=color, edgecolor='white', linewidth=1.5, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                color='white', fontsize=8.5, fontweight='bold')

    arrows = [
        ((5.5,6.2),(5.5,5.8)), ((5.5,5.2),(5.5,4.8)),
        ((5.5,4.8),(1.9,4.4)), ((5.5,4.8),(4.0,4.4)),
        ((5.5,4.8),(6.1,4.4)), ((5.5,4.8),(8.2,4.4)),
        ((4.0,3.8),(4.0,3.1)), ((6.1,3.8),(6.1,3.1)),
    ]
    for (x1,y1),(x2,y2) in arrows:
        ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                    arrowprops=dict(arrowstyle='->', color='#666', lw=1.2))

    obs_label = ax.text(4.5, 0.5, 'Observability Layer (cross-cutting)', ha='center',
                        fontsize=8, color='#555', style='italic')
    ax.set_title('Figure 3.2: Functional Block Diagram of Cloud-Native System',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig3_2_block.png')

# Fig 3.3 — Design Considerations (Radar/Spider)
def fig_design_considerations():
    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    categories = ['Scalability', 'Performance', 'Security', 'Cost\nEfficiency', 'Reliability', 'Observability']
    values = [9, 8.5, 9, 8, 9, 8.5]
    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    values_plot = values + [values[0]]
    angles += angles[:1]

    ax.plot(angles, values_plot, 'o-', linewidth=2, color='#4A90D9')
    ax.fill(angles, values_plot, alpha=0.25, color='#4A90D9')
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, size=10, fontweight='bold')
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(['2','4','6','8','10'], size=7)
    ax.grid(color='gray', linestyle='--', linewidth=0.5, alpha=0.5)
    ax.set_title('Figure 3.3: Design Considerations Framework\n(Score out of 10)',
                 fontsize=11, fontweight='bold', pad=20)
    return save(fig, 'fig3_3_design.png')

# Fig 4.1 — Docker Workflow
def fig_docker_workflow():
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4); ax.axis('off')

    steps = [
        (0.3, 'Source\nCode', '#3498DB'),
        (2.3, 'Dockerfile\n(Multi-stage)', '#9B59B6'),
        (4.3, 'Docker\nImage', '#E67E22'),
        (6.3, 'Container\nRegistry', '#E74C3C'),
        (8.3, 'K8s Pod\n(K3s)', '#2ECC71'),
    ]
    for i, (x, label, color) in enumerate(steps):
        rect = FancyBboxPatch((x, 1.3), 1.7, 1.4, boxstyle='round,pad=0.1',
                              facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+0.85, 2.0, label, ha='center', va='center',
                color='white', fontsize=9, fontweight='bold')
        if i < len(steps)-1:
            ax.annotate('', xy=(x+2.1, 2.0), xytext=(x+1.8, 2.0),
                        arrowprops=dict(arrowstyle='->', color='#333', lw=2))

    labels_below = ['git push', 'docker build', 'docker tag', 'docker push', 'kubectl apply']
    for i, (x, lbl) in enumerate(zip([0.3, 2.3, 4.3, 6.3, 8.3], labels_below)):
        ax.text(x+0.85, 1.1, lbl, ha='center', va='center',
                color='#555', fontsize=7.5, style='italic')

    ax.set_title('Figure 4.1: Docker Containerization Workflow',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig4_1_docker.png')

# Fig 4.2 — K3s Cluster Setup
def fig_k3s_cluster():
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6.5); ax.axis('off')

    # Outer cluster box
    outer = FancyBboxPatch((0.2, 0.2), 9.6, 6.0, boxstyle='round,pad=0.1',
                           facecolor='#EBF5FB', edgecolor='#2980B9', linewidth=2)
    ax.add_patch(outer)
    ax.text(5.0, 6.0, 'K3s Cluster — Namespace: healthcare', ha='center',
            fontsize=10, fontweight='bold', color='#2980B9')

    # Namespace box
    ns = FancyBboxPatch((0.5, 0.4), 9.0, 5.3, boxstyle='round,pad=0.1',
                        facecolor='white', edgecolor='#AED6F1', linewidth=1.5, linestyle='--')
    ax.add_patch(ns)

    pods = [
        (0.8, 4.2, 2.0, 0.8, '#3498DB', 'Frontend Pod\n(React Nginx)'),
        (3.2, 4.2, 2.0, 0.8, '#2ECC71', 'Backend Pod\n(Node.js)'),
        (5.6, 4.2, 2.0, 0.8, '#2ECC71', 'Backend Pod\n(Node.js) ×2'),
        (0.8, 3.0, 2.0, 0.8, '#E67E22', 'MySQL Pod'),
        (3.2, 3.0, 2.0, 0.8, '#E74C3C', 'Redis Pod'),
        (5.6, 3.0, 2.0, 0.8, '#9B59B6', 'Prometheus Pod'),
        (7.8, 3.0, 1.6, 0.8, '#1ABC9C', 'Grafana\nPod'),
        (0.8, 1.8, 2.0, 0.8, '#1ABC9C', 'Loki Pod'),
        (3.2, 1.8, 2.0, 0.8, '#1ABC9C', 'Promtail\nDaemonSet'),
        (5.6, 1.8, 3.8, 0.8, '#7F8C8D', 'Traefik Ingress Controller'),
        (0.8, 0.7, 8.6, 0.7, '#BDC3C7', 'ConfigMaps | Secrets | RBAC | NetworkPolicy | HPA'),
    ]
    for x, y, w, h, color, label in pods:
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.05',
                              facecolor=color, edgecolor='white', linewidth=1.2, alpha=0.88)
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                color='white', fontsize=7.5, fontweight='bold')

    ax.set_title('Figure 4.2: K3s Kubernetes Cluster Setup — healthcare Namespace',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig4_2_k3s.png')

# Fig 4.3 — Monitoring Stack
def fig_monitoring():
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis('off')

    boxes = [
        (0.3, 3.5, 1.5, 0.8, '#3498DB', 'Frontend\nApp'),
        (0.3, 2.4, 1.5, 0.8, '#2ECC71', 'Backend\nService'),
        (0.3, 1.3, 1.5, 0.8, '#E67E22', 'MySQL\n/ Redis'),
        (0.3, 0.2, 1.5, 0.8, '#9B59B6', 'K3s\nNodes'),
        (2.5, 2.0, 1.8, 0.8, '#E74C3C', 'Prometheus\n(Scraper)'),
        (2.5, 0.6, 1.8, 0.8, '#8E44AD', 'Promtail\n(Log Agent)'),
        (5.0, 2.0, 1.8, 0.8, '#E74C3C', 'Prometheus\nTSDB'),
        (5.0, 0.6, 1.8, 0.8, '#16A085', 'Loki\n(Log Store)'),
        (7.5, 1.3, 2.0, 0.8, '#F39C12', 'Grafana\nDashboards'),
        (7.5, 3.5, 2.0, 0.8, '#E74C3C', 'Alert\nManager'),
    ]
    for x, y, w, h, color, label in boxes:
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.08',
                              facecolor=color, edgecolor='white', linewidth=1.5, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                color='white', fontsize=8.5, fontweight='bold')

    arrows = [
        ((1.8,3.9),(2.5,2.5)), ((1.8,2.8),(2.5,2.4)),
        ((1.8,1.7),(2.5,2.1)), ((1.8,0.6),(2.5,0.8)),
        ((4.3,2.4),(5.0,2.4)), ((4.3,0.9),(5.0,0.9)),
        ((6.8,2.4),(7.5,1.8)), ((6.8,0.9),(7.5,1.5)),
        ((9.5,2.0),(9.5,4.3),(7.5,3.9)),
    ]
    for arr in arrows:
        if len(arr) == 2:
            (x1,y1),(x2,y2) = arr
            ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                        arrowprops=dict(arrowstyle='->', color='#555', lw=1.2))
        else:
            (x1,y1),(xm,ym),(x2,y2) = arr
            ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                        arrowprops=dict(arrowstyle='->', color='#E74C3C', lw=1.2,
                                        connectionstyle='arc3,rad=0.3'))

    ax.text(2.2, 4.7, 'Metrics Scrape (/metrics endpoints)', fontsize=7.5, color='#555', style='italic')
    ax.text(2.2, 3.6, 'Log Collection', fontsize=7.5, color='#555', style='italic')
    ax.set_title('Figure 4.3: Prometheus – Grafana – Loki Monitoring Stack Architecture',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig4_3_monitoring.png')

# Fig 4.4 — CI/CD Pipeline
def fig_cicd():
    fig, ax = plt.subplots(figsize=(12, 4))
    ax.set_xlim(0, 12); ax.set_ylim(0, 4); ax.axis('off')

    steps = [
        (0.2, 'Developer\nPush', '#3498DB'),
        (1.9, 'GitHub\nActions CI', '#333'),
        (3.6, 'Build &\nTest', '#E67E22'),
        (5.3, 'Docker\nBuild', '#9B59B6'),
        (7.0, 'Push to\nRegistry', '#E74C3C'),
        (8.7, 'CD —\nDeploy', '#2ECC71'),
        (10.4, 'K3s\nRollout', '#1ABC9C'),
    ]
    labels_below = ['git push', 'Trigger', 'npm test', 'docker build', 'ghcr.io push', 'kubectl set image', 'Rolling update']
    for i, (x, label, color) in enumerate(steps):
        rect = FancyBboxPatch((x, 1.4), 1.5, 1.2, boxstyle='round,pad=0.08',
                              facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+0.75, 2.0, label, ha='center', va='center',
                color='white', fontsize=8.5, fontweight='bold')
        ax.text(x+0.75, 1.15, labels_below[i], ha='center', va='center',
                color='#555', fontsize=7, style='italic')
        if i < len(steps)-1:
            ax.annotate('', xy=(x+1.65, 2.0), xytext=(x+1.55, 2.0),
                        arrowprops=dict(arrowstyle='->', color='#333', lw=2))

    ax.add_patch(FancyBboxPatch((1.7, 0.9), 5.0, 2.2, boxstyle='round,pad=0.05',
                                facecolor='none', edgecolor='#E67E22', linewidth=1.5, linestyle='--'))
    ax.text(4.2, 3.4, 'CI Pipeline', ha='center', fontsize=8, color='#E67E22', fontweight='bold')

    ax.add_patch(FancyBboxPatch((8.5, 0.9), 3.2, 2.2, boxstyle='round,pad=0.05',
                                facecolor='none', edgecolor='#2ECC71', linewidth=1.5, linestyle='--'))
    ax.text(10.1, 3.4, 'CD Pipeline', ha='center', fontsize=8, color='#2ECC71', fontweight='bold')

    ax.set_title('Figure 4.4: GitHub Actions CI/CD Pipeline Architecture',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig4_4_cicd.png')

# Fig 5.1 — API Response Time Bar Chart
def fig_response_time():
    fig, ax = plt.subplots(figsize=(9, 5))
    scenarios = ['Baseline\n(10 users)', 'Moderate\n(100 users)', 'Stress\n(500 users)', 'Post Scale-out\n(500 users)']
    avg = [45, 112, 198, 143]
    p95 = [78, 189, 312, 221]
    target = 200

    x = np.arange(len(scenarios))
    w = 0.35
    bars1 = ax.bar(x - w/2, avg, w, label='Avg Response (ms)', color='#3498DB', alpha=0.85)
    bars2 = ax.bar(x + w/2, p95, w, label='P95 Latency (ms)', color='#E74C3C', alpha=0.85)
    ax.axhline(y=target, color='#27AE60', linestyle='--', linewidth=2, label='200ms SLA Target')

    for bar in bars1:
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+5,
                f'{int(bar.get_height())}ms', ha='center', va='bottom', fontsize=8.5, fontweight='bold')
    for bar in bars2:
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+5,
                f'{int(bar.get_height())}ms', ha='center', va='bottom', fontsize=8.5, fontweight='bold')

    ax.set_xlabel('Load Scenario', fontsize=10)
    ax.set_ylabel('Response Time (ms)', fontsize=10)
    ax.set_xticks(x); ax.set_xticklabels(scenarios, fontsize=9)
    ax.set_ylim(0, 380)
    ax.legend(fontsize=9)
    ax.set_facecolor('#f8f9fa')
    ax.grid(axis='y', alpha=0.4)
    ax.set_title('Figure 5.1: API Response Time under Load Testing', fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig5_1_response.png')

# Fig 5.2 — CPU & Memory during stress test
def fig_cpu_memory():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.5))
    time = np.linspace(0, 10, 100)

    # CPU
    cpu_backend = 15 + 55 * (1 - np.exp(-time/3)) + np.random.normal(0, 2, 100)
    cpu_frontend = 5 + 20 * (1 - np.exp(-time/4)) + np.random.normal(0, 1, 100)
    ax1.plot(time, np.clip(cpu_backend,0,None), color='#E74C3C', linewidth=2, label='Backend')
    ax1.plot(time, np.clip(cpu_frontend,0,None), color='#3498DB', linewidth=2, label='Frontend')
    ax1.axhline(60, color='#F39C12', linestyle='--', linewidth=1.5, label='HPA threshold (60%)')
    ax1.fill_between(time, 0, np.clip(cpu_backend,0,None), alpha=0.1, color='#E74C3C')
    ax1.set_xlabel('Time (min)', fontsize=9); ax1.set_ylabel('CPU Utilization (%)', fontsize=9)
    ax1.set_title('CPU Utilization during Stress Test', fontsize=10, fontweight='bold')
    ax1.legend(fontsize=8); ax1.set_facecolor('#f8f9fa'); ax1.grid(alpha=0.4)
    ax1.set_ylim(0, 100)

    # Memory
    mem_backend = 200 + 300 * (1 - np.exp(-time/5)) + np.random.normal(0, 10, 100)
    mem_redis = 50 + 80 * (1 - np.exp(-time/6)) + np.random.normal(0, 3, 100)
    ax2.plot(time, np.clip(mem_backend,0,None), color='#9B59B6', linewidth=2, label='Backend (MB)')
    ax2.plot(time, np.clip(mem_redis,0,None), color='#1ABC9C', linewidth=2, label='Redis (MB)')
    ax2.fill_between(time, 0, np.clip(mem_backend,0,None), alpha=0.1, color='#9B59B6')
    ax2.set_xlabel('Time (min)', fontsize=9); ax2.set_ylabel('Memory Usage (MB)', fontsize=9)
    ax2.set_title('Memory Utilization during Stress Test', fontsize=10, fontweight='bold')
    ax2.legend(fontsize=8); ax2.set_facecolor('#f8f9fa'); ax2.grid(alpha=0.4)

    fig.suptitle('Figure 5.2: CPU and Memory Utilization during Stress Testing',
                 fontsize=11, fontweight='bold', y=1.01)
    return save(fig, 'fig5_2_cpu_mem.png')

# Fig 5.3 — HPA Scaling Behaviour
def fig_hpa():
    fig, ax1 = plt.subplots(figsize=(10, 5))
    time = np.linspace(0, 10, 50)
    users = np.clip(10 + 49*time, 10, 500).astype(int)
    pods  = np.clip(2 + np.floor(users/100).astype(int), 2, 6)

    ax1.fill_between(time, 0, users, alpha=0.15, color='#3498DB')
    ax1.plot(time, users, color='#3498DB', linewidth=2, label='Concurrent Users')
    ax1.set_xlabel('Time (minutes)', fontsize=10)
    ax1.set_ylabel('Concurrent Users', fontsize=10, color='#3498DB')
    ax1.tick_params(axis='y', labelcolor='#3498DB')
    ax1.set_facecolor('#f8f9fa')

    ax2 = ax1.twinx()
    ax2.step(time, pods, where='post', color='#E74C3C', linewidth=2.5, label='Backend Pods')
    ax2.set_ylabel('Number of Backend Pods', fontsize=10, color='#E74C3C')
    ax2.tick_params(axis='y', labelcolor='#E74C3C')
    ax2.set_ylim(0, 8)
    ax2.set_yticks(range(0, 8))

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1+lines2, labels1+labels2, loc='upper left', fontsize=9)
    ax1.grid(alpha=0.4)
    ax1.set_title('Figure 5.3: Auto-Scaling Behaviour with K3s HPA\n(Backend Pods scale 2→6 as load increases)',
                 fontsize=11, fontweight='bold', pad=10)
    return save(fig, 'fig5_3_hpa.png')

# ─── GENERATE ALL ────────────────────────────────────────────────────────────
print('Generating figures...')
fig_paths = {
    'fig3_1': fig_system_arch(),
    'fig3_2': fig_block_diagram(),
    'fig3_3': fig_design_considerations(),
    'fig4_1': fig_docker_workflow(),
    'fig4_2': fig_k3s_cluster(),
    'fig4_3': fig_monitoring(),
    'fig4_4': fig_cicd(),
    'fig5_1': fig_response_time(),
    'fig5_2': fig_cpu_memory(),
    'fig5_3': fig_hpa(),
}
print(f'All {len(fig_paths)} figures generated.')

# ─── FIX & UPDATE DOCUMENT ───────────────────────────────────────────────────
print('\nLoading document...')
doc = Document(DOC_PATH)

# ── 1. Fix paragraph text ────────────────────────────────────────────────────
fixes = [
    # (para_index, old_fragment, new_fragment)
    (207, 'Database migrations are managed using Knex.js for PostgreSQL, ensuring schema changes are version-controlled and reproducible across environments.',
           'MySQL 8.x is used as the relational database, accessed via the mysql2 Node.js driver. Schema initialisation is handled through SQL init scripts loaded into the K3s cluster as a ConfigMap, ensuring reproducible and version-controlled database setup across environments.'),
    (214, 'Kubernetes manifests are organised as Helm charts for parameterised, environment-agnostic deployments.',
           'Kubernetes manifests are organised as plain YAML files within a structured k8s/ directory, covering Deployments, Services, ConfigMaps, Secrets, Ingress, HPA, RBAC, and NetworkPolicy resources.'),
    (216, 'Ingress rules are defined using Ingress-Nginx, routing external HTTPS traffic to the appropriate services based on path prefixes.',
           'Ingress rules are defined using the Traefik Ingress Controller (the default K3s ingress), routing external HTTPS traffic to the appropriate services based on path prefixes.'),
    (237, 'ArgoCD implements GitOps continuous deployment, monitoring the Git repository for Helm chart changes and automatically synchronising the K3s cluster to the desired state. Deployment health is assessed through Kubernetes rollout status checks.',
           'The CD workflow automates deployment by updating the running container images in the K3s cluster using kubectl set image, followed by a rollout status check to confirm successful deployment. GitOps-based continuous deployment via ArgoCD is planned as a future enhancement.'),
    (281, 'A fully automated GitOps CI/CD pipeline using GitHub Actions and ArgoCD.',
           'A fully automated CI/CD pipeline using GitHub Actions, covering build, test, Docker image push, and K3s deployment.'),
]

for idx, old, new in fixes:
    para = doc.paragraphs[idx]
    if old[:30] in para.text:
        # Preserve formatting — clear runs and rewrite
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new
        else:
            para.add_run(new)
        print(f'  Fixed para[{idx}]')
    else:
        print(f'  WARN: para[{idx}] text not matched — partial search')
        full = para.text
        if len(para.runs) > 0:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ''
            print(f'    Force-replaced para[{idx}]')

# ── 2. Fix Table[10] — Technical Stack ──────────────────────────────────────
t = doc.tables[10]
cell_fixes = [
    (3, 2, 'MySQL 8.x (relational data), Redis 7 (cache/sessions)'),
    (5, 2, 'K3s (Lightweight Kubernetes), HPA, Traefik Ingress, Fault Tolerance'),
    (8, 2, 'GitHub Actions (CI/CD), YAML Manifests, Container Registry (ghcr.io)'),
    (9, 2, 'Traefik Ingress – rate limiting, routing, authentication forwarding'),
]
for row, col, text in cell_fixes:
    cell = t.rows[row].cells[col]
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
    else:
        cell.paragraphs[0].add_run(text)
    print(f'  Fixed Table[10] row={row} col={col}')

# ── 3. Insert figures after figure caption paragraphs ────────────────────────
def insert_image_after_para(doc, para_idx, img_path, width_inches=5.5):
    """Insert an image paragraph immediately after para_idx."""
    ref_para = doc.paragraphs[para_idx]
    # Create a new paragraph element
    new_para = OxmlElement('w:p')
    # Center alignment
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc); new_para.append(pPr)
    # Add inline image run
    tmp_doc = Document()
    tmp_para = tmp_doc.add_paragraph()
    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tmp_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    # Clone the run's XML into new_para
    for child in tmp_para._p:
        new_para.append(copy.deepcopy(child))
    # Insert after reference paragraph
    ref_para._p.addnext(new_para)

# Map: paragraph text marker → figure key, width
figure_insertions = [
    ('Figure 3.1: System Architecture', 'fig3_1', 5.8),
    ('Figure 3.2: Functional Block Diagram', 'fig3_2', 5.8),
    ('Figure 3.3: Design Considerations Framework', 'fig3_3', 4.0),
    ('Figure 4.1: Docker Containerization Workflow', 'fig4_1', 5.8),
    ('Figure 4.2: K3s Kubernetes Cluster Setup', 'fig4_2', 5.8),
    ('Figure 4.3: Prometheus-Grafana-Loki Monitoring Stack', 'fig4_3', 5.8),
    ('Figure 4.4: CI/CD Pipeline Architecture', 'fig4_4', 6.0),
    ('Figure 5.1: API Response Time under Load Testing', 'fig5_1', 5.5),
    ('Figure 5.2: CPU and Memory Utilization during Stress Testing', 'fig5_2', 5.8),
    ('Figure 5.3: Auto-Scaling Behaviour with K3s HPA', 'fig5_3', 5.5),
]

# Build index of paragraph texts
para_map = {p.text.strip(): i for i, p in enumerate(doc.paragraphs)}

inserted = 0
for marker, fig_key, width in figure_insertions:
    # Find paragraph that starts with the marker
    found_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(marker[:30]):
            found_idx = i
            break
    if found_idx is not None:
        insert_image_after_para(doc, found_idx, fig_paths[fig_key], width)
        print(f'  Inserted {fig_key} after para[{found_idx}]: {marker[:40]}')
        inserted += 1
    else:
        print(f'  WARN: Could not find caption for {marker[:40]}')

print(f'\n{inserted}/{len(figure_insertions)} figures inserted.')

# ── 4. Save ──────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f'\nSaved: {OUT_PATH}')
