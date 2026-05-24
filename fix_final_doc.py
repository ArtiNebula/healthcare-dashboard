"""
Fix 2024MT03117-Final-Regenerated.docx:
1. Fill Table of Contents with page numbers
2. Fill List of Abbreviations
3. Fill List of Figures
4. Fill List of Tables
5. Insert REAL figures (screenshots + diagrams) at appropriate positions with proper captions
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\2024MT03117-Final-Regenerated.docx'
OUT_PATH = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\2024MT03117-Final-v3.docx'
IMGS     = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\dissertation_imgs'
SCRN     = r'd:\M.tech\SEM IV\Healthcare Dashboard Development\screenshots'

doc = Document(DOC_PATH)

# ── helpers ───────────────────────────────────────────────────────────────────
def para_after(anchor_elem, style_name='Normal'):
    new_p = OxmlElement('w:p')
    anchor_elem.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            try:
                p.style = doc.styles[style_name]
            except Exception:
                pass
            return p
    return None

def center(p):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

def set_spacing(p, before=0, after=6):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    spc = pPr.find(qn('w:spacing'))
    if spc is None:
        spc = OxmlElement('w:spacing')
        pPr.append(spc)
    spc.set(qn('w:before'), str(before))
    spc.set(qn('w:after'),  str(after))

def add_toc_line(anchor_elem, left_text, page_no, bold=False, indent_cm=0):
    """Insert a TOC line with dot-leader tab, always after the same anchor."""
    p = para_after(anchor_elem, 'Normal')
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(indent_cm * 567)))
    pPr.append(ind)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)
    p._element.insert(0, pPr)
    r1 = p.add_run(left_text)
    r1.bold = bold
    r1.font.size = Pt(11)
    r2 = p.add_run('\t' + str(page_no))
    r2.bold = bold
    r2.font.size = Pt(11)
    return p._element

def insert_figure(anchor_elem, img_path, caption_text, width_in=5.8):
    """Insert image then caption after anchor_elem. Returns caption element."""
    # 1. blank spacer before image
    sp = para_after(anchor_elem, 'Normal')
    set_spacing(sp, before=60, after=0)

    # 2. image paragraph
    img_p = para_after(sp._element, 'Normal')
    center(img_p)
    set_spacing(img_p, before=0, after=0)
    if os.path.exists(img_path):
        try:
            run = img_p.add_run()
            run.add_picture(img_path, width=Inches(width_in))
        except Exception as e:
            img_p.add_run(f'[Image error: {e}]')
    else:
        img_p.add_run(f'[Missing: {os.path.basename(img_path)}]')

    # 3. caption paragraph
    cap_p = para_after(img_p._element, 'Normal')
    center(cap_p)
    set_spacing(cap_p, before=4, after=80)
    run = cap_p.add_run(caption_text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    return cap_p._element

def find_section_end(heading_text):
    """Return last non-empty para element before the next Heading after heading_text."""
    found = False
    heading_elem = None
    last_elem = None
    for p in doc.paragraphs:
        if not found:
            if p.text.strip() == heading_text:
                found = True
                heading_elem = p._element
                last_elem = p._element
        else:
            if p.style.name.startswith('Heading'):
                return last_elem if (last_elem is not None and last_elem is not heading_elem) else heading_elem
            if p.text.strip():
                last_elem = p._element
    return last_elem if last_elem is not None else heading_elem

# ── page number map ───────────────────────────────────────────────────────────
P = {
    'cert':'iii', 'ack':'iv', 'abs':'v',
    'abbrev':'ix', 'lof':'x', 'lot':'xi',
    'ch1':1,  '1.1':1,  '1.2':2,  '1.3':3,  '1.4':4,
    'ch2':7,  '2.1':7,  '2.2':8,  '2.3':10, '2.4':11, '2.5':12, '2.6':12,
    'ch3':14, '3.1':14, '3.2':15, '3.3':15, '3.4':17,
    'ch4':19, '4.1':19, '4.2':21, '4.3':22, '4.4':23,
              '4.5':24, '4.6':25, '4.7':26, '4.7.1':26,'4.7.2':27,
    'ch5':30, '5.1':30, '5.2':31, '5.3':32, '5.4':33, '5.5':34,
    'ch6':36, '6.1':36, '6.2':37, '6.3':38, '6.4':39, '6.5':40,
    'ch7':42, '7.1':42, '7.2':43, '7.3':44,
    'ch8':46, '8.1':46, '8.2':48, '8.3':49, '8.4':50,
    'ch9':52, '9.1':52, '9.2':55, '9.3':56, '9.4':57,
    'ch10':60,'10.1':60,'10.2':62,
    'glos':65,'appA':66,'appB':69,'appD':71,'ref':73,
}

# ── figure & table metadata ───────────────────────────────────────────────────
FIGURES = [
    ('Figure 3.1', 'High-Level System Architecture',            'fig3_1_arch.png',    IMGS, P['3.3']),
    ('Figure 3.2', 'System Component Block Diagram',            'fig3_2_block.png',   IMGS, P['3.4']),
    ('Figure 4.1', 'User Dashboard Interface',                  '02_user_dashboard.png', SCRN, P['4.2']),
    ('Figure 4.2', 'Symptom Entry Interface',                   '03_add_symptoms.png',   SCRN, P['4.3']),
    ('Figure 4.3', 'AI Analysis and Suggestions Interface',     '05_ai_suggestions.png', SCRN, P['4.5']),
    ('Figure 4.4', 'Admin Dashboard Interface',                 '06_admin_dashboard.png',SCRN, P['4.7.1']),
    ('Figure 4.5', 'User Monitoring Interface',                 '07_user_monitoring.png',SCRN, P['4.7.2']),
    ('Figure 6.1', 'Docker Multi-Stage Build and Deployment',   'fig4_1_docker.png',  IMGS, P['6.1']),
    ('Figure 6.2', 'K3s Cluster Architecture and HPA Scaling',  'fig4_2_k3s.png',     IMGS, P['6.4']),
    ('Figure 7.1', 'Prometheus Custom Metrics Dashboard',       '12_prometheus.png',   SCRN, P['7.1']),
    ('Figure 7.2', 'Grafana Monitoring Dashboard Overview',     '11_grafana.png',      SCRN, P['7.2']),
    ('Figure 9.1', 'API Response Time Distribution (k6 Test)',  'fig5_1_response.png', IMGS, P['9.1']),
    ('Figure 9.2', 'CPU and Memory Usage Under Load',           'fig5_2_cpu_mem.png',  IMGS, P['9.4']),
    ('Figure 9.3', 'HPA Scaling Behaviour Under Traffic Spike', 'fig5_3_hpa.png',      IMGS, P['9.3']),
]

TABLES = [
    ('Table 3.1', 'Functional Requirements',                               P['3.1']),
    ('Table 3.2', 'Non-Functional Requirements',                           P['3.2']),
    ('Table 3.3', 'Technology Stack with Justification',                   P['3.4']),
    ('Table 5.1', 'REST API Endpoint Summary',                             P['5.1']),
    ('Table 6.1', 'Docker Image Sizes: Before and After Multi-Stage Build',P['6.1']),
    ('Table 6.2', 'Kubernetes Resource Requests and Limits',               P['6.4']),
    ('Table 7.1', 'Prometheus Custom Metrics',                             P['7.1']),
    ('Table 8.1', 'HIPAA Security Rule Safeguard Mapping',                 P['8.1']),
    ('Table 9.1', 'Load Test Summary (k6)',                                P['9.1']),
    ('Table 9.2', 'Security Vulnerability Scan Results (OWASP ZAP)',       P['9.3']),
]

ABBREVIATIONS = [
    ('AI',    'Artificial Intelligence'),
    ('API',   'Application Programming Interface'),
    ('BITS',  'Birla Institute of Technology and Science'),
    ('CI/CD', 'Continuous Integration / Continuous Deployment'),
    ('CPU',   'Central Processing Unit'),
    ('CRUD',  'Create, Read, Update, Delete'),
    ('CSS',   'Cascading Style Sheets'),
    ('DNS',   'Domain Name System'),
    ('GDPR',  'General Data Protection Regulation'),
    ('HPA',   'Horizontal Pod Autoscaler'),
    ('HIPAA', 'Health Insurance Portability and Accountability Act'),
    ('HTML',  'HyperText Markup Language'),
    ('HTTP',  'HyperText Transfer Protocol'),
    ('HTTPS', 'HyperText Transfer Protocol Secure'),
    ('JSON',  'JavaScript Object Notation'),
    ('JWT',   'JSON Web Token'),
    ('K3s',   'Lightweight Kubernetes Distribution'),
    ('K8s',   'Kubernetes'),
    ('ML',    'Machine Learning'),
    ('MySQL', 'My Structured Query Language'),
    ('Nginx', 'Engine-X Web Server / Reverse Proxy'),
    ('ORM',   'Object-Relational Mapping'),
    ('OWASP', 'Open Web Application Security Project'),
    ('RBAC',  'Role-Based Access Control'),
    ('REST',  'Representational State Transfer'),
    ('SPA',   'Single Page Application'),
    ('SQL',   'Structured Query Language'),
    ('SSL',   'Secure Sockets Layer'),
    ('TLS',   'Transport Layer Security'),
    ('TS',    'TypeScript'),
    ('TTL',   'Time To Live'),
    ('UI',    'User Interface'),
    ('UX',    'User Experience'),
    ('WILP',  'Work Integrated Learning Programme'),
    ('YAML',  "YAML Ain't Markup Language"),
]

TOC_ENTRIES = [
    ('CERTIFICATE',                                     P['cert'], True, 0),
    ('ACKNOWLEDGEMENTS',                                P['ack'],  True, 0),
    ('ABSTRACT',                                        P['abs'],  True, 0),
    ('LIST OF ABBREVIATIONS',                           P['abbrev'],True,0),
    ('LIST OF FIGURES',                                 P['lof'],  True, 0),
    ('LIST OF TABLES',                                  P['lot'],  True, 0),
    ('CHAPTER 1 – INTRODUCTION',                        P['ch1'],  True, 0),
    ('1.1  Background and Motivation',                  P['1.1'],  False,0.8),
    ('1.2  Problem Statement',                          P['1.2'],  False,0.8),
    ('1.3  Objectives',                                 P['1.3'],  False,0.8),
    ('1.4  Scope and Limitations',                      P['1.4'],  False,0.8),
    ('CHAPTER 2 – LITERATURE REVIEW',                   P['ch2'],  True, 0),
    ('2.1  Digital Health Platforms',                   P['2.1'],  False,0.8),
    ('2.2  Cloud-Native Application Architectures',     P['2.2'],  False,0.8),
    ('2.3  Containerisation and Orchestration',         P['2.3'],  False,0.8),
    ('2.4  AI in Symptom Analysis',                     P['2.4'],  False,0.8),
    ('2.5  Observability and Monitoring',               P['2.5'],  False,0.8),
    ('2.6  Healthcare Data Security and HIPAA',         P['2.6'],  False,0.8),
    ('CHAPTER 3 – SYSTEM DESIGN AND ARCHITECTURE',      P['ch3'],  True, 0),
    ('3.1  Functional Requirements',                    P['3.1'],  False,0.8),
    ('3.2  Non-Functional Requirements',                P['3.2'],  False,0.8),
    ('3.3  High-Level Architecture',                    P['3.3'],  False,0.8),
    ('3.4  Technology Stack Justification',             P['3.4'],  False,0.8),
    ('CHAPTER 4: FRONTEND IMPLEMENTATION',              P['ch4'],  True, 0),
    ('4.1  Application Structure and Routing',          P['4.1'],  False,0.8),
    ('4.2  User Dashboard',                             P['4.2'],  False,0.8),
    ('4.3  Symptom Entry Module',                       P['4.3'],  False,0.8),
    ('4.4  Health History Module',                      P['4.4'],  False,0.8),
    ('4.5  AI Analysis Page',                           P['4.5'],  False,0.8),
    ('4.6  User Profile Module',                        P['4.6'],  False,0.8),
    ('4.7  Administrative Interface',                   P['4.7'],  False,0.8),
    ('    4.7.1  Admin Dashboard',                      P['4.7.1'],False,1.6),
    ('    4.7.2  User Monitoring',                      P['4.7.2'],False,1.6),
    ('CHAPTER 5: BACKEND IMPLEMENTATION',               P['ch5'],  True, 0),
    ('5.1  API Architecture',                           P['5.1'],  False,0.8),
    ('5.2  Authentication and Session Management',      P['5.2'],  False,0.8),
    ('5.3  Data Models',                                P['5.3'],  False,0.8),
    ('5.4  AI Analysis Engine',                         P['5.4'],  False,0.8),
    ('5.5  Rate Limiting and Security Middleware',      P['5.5'],  False,0.8),
    ('CHAPTER 6: CLOUD-NATIVE DEPLOYMENT',              P['ch6'],  True, 0),
    ('6.1  Containerisation with Docker',               P['6.1'],  False,0.8),
    ('6.2  Docker Compose for Local Orchestration',     P['6.2'],  False,0.8),
    ('6.3  K3s Kubernetes Cluster',                     P['6.3'],  False,0.8),
    ('6.4  Horizontal Pod Autoscaling',                 P['6.4'],  False,0.8),
    ('6.5  Service Discovery and Networking',           P['6.5'],  False,0.8),
    ('CHAPTER 7: OBSERVABILITY AND MONITORING',         P['ch7'],  True, 0),
    ('7.1  Prometheus Metrics Collection',              P['7.1'],  False,0.8),
    ('7.2  Grafana Dashboards',                         P['7.2'],  False,0.8),
    ('7.3  Alerting Strategy',                          P['7.3'],  False,0.8),
    ('CHAPTER 8: SECURITY AND COMPLIANCE',              P['ch8'],  True, 0),
    ('8.1  Encryption at Rest and in Transit',          P['8.1'],  False,0.8),
    ('8.2  Role-Based Access Control',                  P['8.2'],  False,0.8),
    ('8.3  Input Validation and Sanitisation',          P['8.3'],  False,0.8),
    ('8.4  Vulnerability Assessment',                   P['8.4'],  False,0.8),
    ('CHAPTER 9: TESTING, RESULTS, AND EVALUATION',    P['ch9'],  True, 0),
    ('9.1  Unit and Integration Testing',               P['9.1'],  False,0.8),
    ('9.2  Performance Testing',                        P['9.2'],  False,0.8),
    ('9.3  Security Testing',                           P['9.3'],  False,0.8),
    ('9.4  System Evaluation and Discussion',           P['9.4'],  False,0.8),
    ('CHAPTER 10: CONCLUSIONS AND FUTURE WORK',        P['ch10'], True, 0),
    ('10.1  Conclusions',                               P['10.1'], False,0.8),
    ('10.2  Future Work',                               P['10.2'], False,0.8),
    ('GLOSSARY',                                        P['glos'], True, 0),
    ('APPENDIX A: Project Directory Structure',         P['appA'], True, 0),
    ('APPENDIX B: API Endpoint Reference',              P['appB'], True, 0),
    ('APPENDIX D: PPT Presentation Outline',            P['appD'], True, 0),
    ('REFERENCES',                                      P['ref'],  True, 0),
]

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
print("Filling Table of Contents ...")
toc_anchor = doc.paragraphs[107]._element   # empty para right after "TABLE OF CONTENTS"
for text, page, bold, indent in reversed(TOC_ENTRIES):
    add_toc_line(toc_anchor, text, page, bold, indent)

# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: LIST OF ABBREVIATIONS
# ─────────────────────────────────────────────────────────────────────────────
print("Filling List of Abbreviations ...")
abbrev_anchor = None
for p in doc.paragraphs:
    if p.text.strip() == 'LIST OF ABBREVIATIONS':
        abbrev_anchor = p._element.getnext()
        break
if abbrev_anchor is not None:
    for abbr, full in reversed(ABBREVIATIONS):
        p = para_after(abbrev_anchor, 'Normal')
        r1 = p.add_run(f'{abbr:<8}')
        r1.bold = True; r1.font.size = Pt(11)
        r2 = p.add_run(f'  {full}')
        r2.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: LIST OF FIGURES
# ─────────────────────────────────────────────────────────────────────────────
print("Filling List of Figures ...")
lof_anchor = None
for p in doc.paragraphs:
    if p.text.strip() == 'LIST OF FIGURES':
        lof_anchor = p._element.getnext()
        break
if lof_anchor is not None:
    for fig_id, fig_title, _f, _d, pg in reversed(FIGURES):
        add_toc_line(lof_anchor, f'{fig_id}:  {fig_title}', pg, False, 0)

# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: LIST OF TABLES
# ─────────────────────────────────────────────────────────────────────────────
print("Filling List of Tables ...")
lot_anchor = None
for p in doc.paragraphs:
    if p.text.strip() == 'LIST OF TABLES':
        lot_anchor = p._element.getnext()
        break
if lot_anchor is not None:
    for tbl_id, tbl_title, pg in reversed(TABLES):
        add_toc_line(lot_anchor, f'{tbl_id}:  {tbl_title}', pg, False, 0)

# ─────────────────────────────────────────────────────────────────────────────
# STEP 5: INSERT FIGURES INTO DOCUMENT BODY  (bottom → top)
# ─────────────────────────────────────────────────────────────────────────────
print("Inserting figures ...")

# Maps: section heading text → (figure label, filename, dir, width_inches)
BODY_FIGURES = [
    # bottom of document first
    ('9.4 System Evaluation and Discussion',
        'Figure 9.2: CPU and Memory Usage Under Load',      'fig5_2_cpu_mem.png', IMGS, 5.2),
    ('9.3 Security Testing',
        'Figure 9.3: HPA Scaling Behaviour Under Traffic Spike', 'fig5_3_hpa.png', IMGS, 5.2),
    ('9.1 Unit and Integration Testing',
        'Figure 9.1: API Response Time Distribution (k6 Test)', 'fig5_1_response.png', IMGS, 5.2),
    ('7.2 Grafana Dashboards',
        'Figure 7.2: Grafana Monitoring Dashboard Overview', '11_grafana.png',    SCRN, 5.8),
    ('7.1 Prometheus Metrics Collection',
        'Figure 7.1: Prometheus Custom Metrics Dashboard',  '12_prometheus.png',  SCRN, 5.5),
    ('6.4 Horizontal Pod Autoscaling',
        'Figure 6.2: K3s Cluster Architecture and HPA Scaling', 'fig4_2_k3s.png', IMGS, 5.2),
    ('6.1 Containerisation with Docker',
        'Figure 6.1: Docker Multi-Stage Build and Deployment', 'fig4_1_docker.png', IMGS, 5.2),
    ('4.7.2 User Monitoring',
        'Figure 4.5: User Monitoring Interface',            '07_user_monitoring.png', SCRN, 5.8),
    ('4.7.1 Admin Dashboard',
        'Figure 4.4: Admin Dashboard Interface',            '06_admin_dashboard.png', SCRN, 5.8),
    ('4.7 Administrative Interface',
        'Figure 4.3: AI Analysis and Suggestions Interface','05_ai_suggestions.png',  SCRN, 5.8),
    ('4.3 Symptom Entry Module',
        'Figure 4.2: Symptom Entry Interface',              '03_add_symptoms.png',    SCRN, 5.8),
    ('4.1 Application Structure and Routing',
        'Figure 4.1: User Dashboard Interface',             '02_user_dashboard.png',  SCRN, 5.8),
    ('3.4 Technology Stack Justification',
        'Figure 3.2: System Component Block Diagram',       'fig3_2_block.png',   IMGS, 5.0),
    ('3.3 High-Level Architecture',
        'Figure 3.1: High-Level System Architecture',       'fig3_1_arch.png',    IMGS, 5.2),
]

for heading, caption, fname, imgdir, width in BODY_FIGURES:
    anchor = find_section_end(heading)
    if anchor is None:
        print(f'  WARNING: section "{heading}" not found')
        continue
    img_path = os.path.join(imgdir, fname)
    insert_figure(anchor, img_path, caption, width_in=width)
    print(f'  ✓ {caption[:55]}')

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
sz = os.path.getsize(OUT_PATH) // 1024
print(f'\nSaved: {OUT_PATH}  ({sz} KB)')
